import streamlit as st
from docx import Document
import re
import io
import time
import os

# --- 1. KONFIGURASI HALAMAN (WAJIB PALING ATAS) ---
st.set_page_config(
    page_title="CBT Converter",
    page_icon="🏫",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# --- 2. DESAIN TAMPILAN (CSS INJECTION) ---
# Ini adalah bagian "Makeover" tampilannya
st.markdown("""
<style>
    /* A. Mengatur Background Halaman (Gradasi Biru) */
    .stApp {
        background: linear-gradient(to bottom right, #E3F2FD, #FFFFFF);
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }

    /* B. Mengatur Container Utama (Efek Kartu Putih) */
    .block-container {
        background-color: rgba(255, 255, 255, 0.95); /* Putih transparan */
        padding: 3rem !important;
        border-radius: 20px;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.1); /* Bayangan halus */
        margin-top: 20px;
        max-width: 900px;
    }

    /* C. Judul Utama */
    .main-title {
        font-size: 2.2rem;
        color: #1A5276; /* Biru Tua */
        text-align: center;
        font-weight: 800;
        margin-bottom: 5px;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    .sub-title {
        font-size: 1.1rem;
        color: #5D6D7E;
        text-align: center;
        margin-bottom: 25px;
        font-weight: 500;
    }

    /* D. Desain Tombol Upload & Download */
    .stButton>button {
        width: 100%;
        border-radius: 10px;
        height: 3.5em;
        background: linear-gradient(90deg, #2980B9 0%, #3498DB 100%);
        color: white;
        font-weight: bold;
        border: none;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: translateY(-2px); /* Efek naik saat ditunjuk */
        box-shadow: 0 6px 8px rgba(0,0,0,0.2);
    }
    
    /* Tombol Download Khusus (Hijau) */
    .stDownloadButton>button {
        width: 100%;
        border-radius: 10px;
        height: 3.5em;
        background: linear-gradient(90deg, #27AE60 0%, #2ECC71 100%);
        color: white;
        font-weight: bold;
        border: none;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .stDownloadButton>button:hover {
        transform: scale(1.02);
    }

    /* E. Desain Kotak Panduan */
    .panduan-box {
        background-color: #F8F9F9;
        padding: 20px;
        border-radius: 15px;
        border-left: 6px solid #F1C40F; /* Garis Kuning */
        margin-top: 20px;
    }

    /* F. Footer */
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #1A5276;
        color: white;
        text-align: center;
        padding: 10px;
        font-size: 12px;
        z-index: 999;
    }
</style>
""", unsafe_allow_html=True)

# --- 3. FUNGSI LOGIKA (BACKEND) ---
def parse_raw_document(file):
    doc = Document(file)
    questions = []
    current_q = {}
    
    re_opt = re.compile(r'^([A-Ea-e])\.\s*(.*)')
    re_num = re.compile(r'^(\d+)\.\s*(.*)')      
    re_key = re.compile(r'(?i)(?:Kunci|Jawaban)\s*[:=]\s*([A-Ea-e])')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue

        key_match = re_key.search(text)
        if key_match and current_q:
            current_q['kunci'] = key_match.group(1).upper()
            continue 

        num_match = re_num.match(text)
        if num_match:
            if current_q: questions.append(current_q)
            current_q = {
                'no': num_match.group(1),
                'text': num_match.group(2),
                'options': {},
                'kunci': 'A'
            }
        
        opt_match = re_opt.match(text)
        if opt_match and current_q:
            opt_label = opt_match.group(1).upper() 
            content = opt_match.group(2)
            current_q['options'][opt_label] = content
            
    if current_q: questions.append(current_q)
    return questions

def create_cbt_doc(questions):
    doc = Document()
    for q in questions:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        data_rows = [("TS", "PG"), ("KD", "1.0.1"), ("KJ", q.get('kunci', 'A').upper()), ("ABS", "")]
        for label, val in data_rows:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = val
        row = table.add_row().cells
        row[0].text = q['no'] + "."
        row[1].text = q['text']
        for opt in ['A', 'B', 'C', 'D', 'E']:
            row = table.add_row().cells
            row[0].text = opt
            row[1].text = q['options'].get(opt, "")
        doc.add_paragraph()
    return doc

# --- 4. TAMPILAN HALAMAN (FRONTEND) ---

# Header dengan Logo
col1, col2, col3 = st.columns([1.2, 4, 1.2])

with col1:
    if os.path.exists("logo_kiri.jpg"):
        st.image("logo_kiri.jpg", use_container_width=True)
    else:
        st.warning("Upload Logo Kiri")

with col2:
    st.markdown('<div class="main-title">Aplikasi Konverter<br>Soal Ujian Online</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-title">SMK NEGERI 1 KOTABUMI</div>', unsafe_allow_html=True)

with col3:
    if os.path.exists("logo_kanan.jpg"):
        st.image("logo_kanan.jpg", use_container_width=True)
    else:
        st.warning("Upload Logo Kanan")

st.markdown("---")

# Tab Navigasi
tab1, tab2 = st.tabs(["🚀 **MULAI KONVERSI**", "📘 **PANDUAN & BANTUAN**"])

# --- ISI TAB 1 ---
with tab1:
    st.info("💡 **Tips:** Pastikan file Word Anda bersih dari format Numbering otomatis.")
    
    uploaded_file = st.file_uploader(
        "📂 Klik disini untuk upload file Soal (.docx)", 
        type=['docx', 'doc']
    )

    if uploaded_file:
        file_name = uploaded_file.name.lower()
        if file_name.endswith('.doc'):
            st.error("⚠️ Format .doc (Word 97-2003) tidak didukung.")
            st.warning("Silakan 'Save As' file Anda menjadi .docx terlebih dahulu.")
        elif file_name.endswith('.docx'):
            with st.status("⚙️ Sedang menganalisis dokumen...", expanded=True) as status:
                time.sleep(0.8) # Efek delay biar kelihatan mikir
                try:
                    data_soal = parse_raw_document(uploaded_file)
                    if not data_soal:
                        status.update(label="Gagal!", state="error")
                        st.error("❌ Tidak ditemukan pola soal yang valid.")
                    else:
                        status.update(label="Analisis Selesai!", state="complete")
                        st.success(f"✅ Berhasil memproses **{len(data_soal)}** soal.")
                        
                        col_a, col_b = st.columns([1, 1])
                        with col_a:
                            st.metric("Jumlah Soal", len(data_soal))
                        with col_b:
                            st.metric("Status Kunci", "Auto-Detect")
                        
                        st.write("") # Spasi
                        doc_output = create_cbt_doc(data_soal)
                        bio = io.BytesIO()
                        doc_output.save(bio)
                        
                        st.download_button(
                            label="📥 DOWNLOAD HASIL KONVERSI (.DOCX)",
                            data=bio.getvalue(),
                            file_name="HASIL_CBT_SMK.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Error: {e}")

# --- ISI TAB 2: PANDUAN LENGKAP ---
with tab2:
    st.markdown('<div class="panduan-box">', unsafe_allow_html=True)
    
    st.markdown("## 📚 Panduan Penggunaan Aplikasi")
    st.write("Ikuti langkah-langkah berikut agar soal dapat terkonversi dengan sempurna.")

    # Bagian 1: Persiapan File (Paling Penting)
    st.info("### 1️⃣ Tahap Persiapan File Word (Wajib Baca)")
    st.write("Sebelum upload, pastikan file Microsoft Word Anda memenuhi syarat berikut:")
    
    col_guide1, col_guide2 = st.columns(2)
    
    with col_guide1:
        st.success("**✅ FORMAT YANG BENAR**")
        st.markdown("""
        **1. Penomoran Manual:**
        Ketik angka dan titik secara manual.
        > `1. Siapakah presiden pertama...`
        
        **2. Pilihan Ganda:**
        Gunakan huruf (A-E) diikuti titik.
        > `A. Soekarno`
        
        **3. Kunci Jawaban:**
        Letakkan di bawah pilihan terakhir.
        > `Kunci: A`
        """)
        
    with col_guide2:
        st.error("**❌ HINDARI HAL INI**")
        st.markdown("""
        **1. Fitur Numbering Otomatis:**
        Jangan gunakan tombol *Numbering* di toolbar Word. Sistem tidak bisa membacanya.
        
        **2. Tanda Kurung:**
        Jangan gunakan tanda kurung untuk nomor/opsi.
        > `1) Siapakah...` (Salah)
        > `A) Soekarno` (Salah)
        """)

    st.markdown("---")

    # Bagian 2: Cara Membersihkan Format (Tips Pro)
    st.warning("### 💡 Tips: Cara Mematikan Numbering Otomatis")
    st.markdown("""
    Jika file soal Anda sudah terlanjur menggunakan *Numbering Otomatis* (saat diklik angkanya ikut terblok semua), lakukan ini:
    1. Tekan **CTRL + A** (Select All) di Word.
    2. Klik ikon **'Clear Formatting'** (Gambar huruf A dengan penghapus merah muda) di menu Home.
    3. Atau, ubah format menjadi **'Normal'**.
    4. Beri nomor ulang secara manual (ketik `1.` spasi `Soal...`).
    """)

    st.markdown("---")

    # Bagian 3: Langkah Konversi
    st.markdown("### 2️⃣ Langkah Konversi di Aplikasi")
    st.markdown("""
    1. Klik tab **🚀 MULAI KONVERSI**.
    2. Klik tombol **"Browse files"** dan pilih file `.docx` Anda.
    3. Tunggu hingga muncul tulisan hijau **"Analisis Selesai!"**.
    4. Periksa jumlah soal yang terdeteksi. Jika sesuai, klik tombol hijau **"DOWNLOAD HASIL"**.
    5. File hasil konversi akan otomatis tersimpan di folder *Downloads* komputer Anda.
    """)

    # Bagian 4: Masalah Umum
    with st.expander("❓ Masalah Umum & Solusinya (Troubleshooting)"):
        st.markdown("""
        | Masalah | Penyebab | Solusi |
        | :--- | :--- | :--- |
        | **Error "File Type"** | Anda mengupload file `.doc` (Word 2003) | Buka file, pilih **Save As**, ganti format ke **.docx**. |
        | **Jumlah Soal 0** | Format nomor tidak terbaca | Cek apakah nomor soal menggunakan fitur *Auto-Numbering*. Hapus dan ketik manual. |
        | **Kunci Jawaban Salah** | Salah ketik kata kunci | Pastikan tertulis `Kunci:` atau `Jawaban:` (pakai titik dua). |
        | **Tabel Berantakan** | Ada gambar/tabel di soal | Hapus gambar/tabel di file mentah, masukkan manual setelah jadi format CBT. |
        """)
    
    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("""
<div class="footer">
    Developed by Gabut Proktor | © 2026
</div>
""", unsafe_allow_html=True)